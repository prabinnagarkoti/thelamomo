import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

doc.add_heading('BizMenu Builder: High Distinction (HD) Revisions', 0)

doc.add_paragraph("Based on the rubric, here are the detailed, updated sections you need to insert into your assignment completely replacing or adding to the current ones to ensure you meet the 85-100% HD criteria.")

# 1. Background & Literature Review Update
doc.add_heading('1. Project Background & Literature Review Revisions', level=1)
doc.add_paragraph('Insert this at the end of section 2.3 (Technologies Selected and Justification):')
p = doc.add_paragraph()
p.add_run("Justification for Next.js App Router: ").bold = True
p.add_run("While earlier iterations of this project plan considered a decoupled React.js, Express, and Node.js architecture, the final implementation pivoted to Next.js (App Router). This shift was strategically justified to leverage serverless deployment capabilities directly on Vercel. For small independent food businesses, infrastructure costs and maintenance are significant barriers. Next.js eliminates the need for a persistent backend server, significantly reducing hosting costs while inherently improving Search Engine Optimization (SEO) through Server-Side Rendering (SSR). This directly addresses the literature's concern regarding the operational burden on small enterprises.")

doc.add_paragraph('Insert this into 2.4 (Gap Identified):')
p = doc.add_paragraph()
p.add_run("Comparative Advantage over Existing White-Label Solutions: ").bold = True
p.add_run("While platforms like GloriaFood and ChowNow offer white-label interfaces, they still bind the restaurant to monthly subscription models and lock customer data within their respective proprietary ecosystems. BizMenu Builder entirely decentralizes this model. Because it is deployed independently via Next.js and MongoDB Atlas, the restaurant retains 100% ownership of their data schema, customer analytics, and transaction routing, free from recurring platform dependency.")

# 2. System Design Updates
doc.add_heading('2. System Design Clarification', level=1)
doc.add_paragraph('Insert this directly under Figure 3.1.1 (Architecture Diagram) and Figure 3.2.1 (ERD Database Model):')
p = doc.add_paragraph()
p.add_run("Database Implementation Note for Prototype: ").bold = True
p.add_run("The Entity Relationship Diagram (ERD) in Figure 3.2.1 illustrates the comprehensive architecture envisioned for the final production build, which includes advanced features like Profiles, Notifications, and Payments. However, in alignment with agile iterative development, the current mid-point prototype specifically implements the core foundation: MenuItem, Order, RestaurantConfig, and User schemas. This scoped approach ensures core ordering pipeline stability before scaling to peripheral notification features in subsequent sprints.")

# 4. Product Backlog & Sprints
doc.add_heading('4. Enhanced Product Backlog (With Story Points & Assignees)', level=1)
doc.add_paragraph("Replace your Table 1 (Product Backlog) with this enhanced matrix:")

table = doc.add_table(rows=1, cols=6)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Epic'
hdr_cells[1].text = 'User Story'
hdr_cells[2].text = 'Acceptance Criteria'
hdr_cells[3].text = 'Priority'
hdr_cells[4].text = 'Story Pts.'
hdr_cells[5].text = 'Status'

stories = [
    ("EP-03", "As a developer, I want to connect the app to MongoDB Atlas...", "Connection established; schema created", "High", "3", "Done"),
    ("EP-01", "As an owner, I want to add menu items...", "Owner inputs data; appears on menu", "High", "5", "Done"),
    ("EP-01", "As an owner, I want to edit and delete items...", "Modified items update instantly", "High", "5", "In Progress"),
    ("EP-02", "As a customer, I want to filter by category...", "Filters display correctly", "High", "2", "Done"),
    ("EP-02", "As a customer, I want to search items by name...", "Search input filters in real-time", "Medium", "3", "Done"),
    ("EP-02", "As a customer, I want to add items to cart...", "Cart updates; order submitted to DB", "High", "8", "In Progress"),
    ("EP-03", "As a developer, I want to deploy to Vercel...", "App builds successfully; accessible via URL", "High", "3", "In Progress")
]

for epic, us, ac, prio, sp, status in stories:
    row_cells = table.add_row().cells
    row_cells[0].text = epic
    row_cells[1].text = us
    row_cells[2].text = ac
    row_cells[3].text = prio
    row_cells[4].text = sp
    row_cells[5].text = status

doc.add_paragraph()
doc.add_paragraph('Insert this into your Sprint Review Minutes:')
p = doc.add_paragraph()
p.add_run("Sprint 1 Retrospective: \n").bold = True
p.add_run("What went well: The GitHub repository, Next.js foundation, and MongoDB connection were established ahead of schedule. \nWhat didn't go well: Mongoose import path resolutions in the Next.js App Router caused unexpected delays in local build testing. \nAction Items for Next Sprint: Implement stricter PR review policies for dependency updates and allocate more paired-programming time for deployment testing.")

p = doc.add_paragraph()
p.add_run("Sprint 2 Retrospective: \n").bold = True
p.add_run("What went well: The frontend deployment succeeded, and the customer menu interface is functioning well with live Next.js rendering.\nWhat didn't go well: Custom domain DNS propagation failed, leading to invalid configuration states on Vercel.\nAction Items for Next Sprint: Pivot focus to resolving DNS 'A' record pointing, and transition immediately into finalized cart state management.")

# 5. Risk Management Updates
doc.add_heading('5. Upgraded Risk Management Matrix', level=1)
doc.add_paragraph("Replace your current Risk Management table with this upgraded version that tracks Likelihood and Sprint Status:")

rtable = doc.add_table(rows=1, cols=5)
rtable.style = 'Table Grid'
hdr_cells = rtable.rows[0].cells
hdr_cells[0].text = 'Potential Risk'
hdr_cells[1].text = 'Impact'
hdr_cells[2].text = 'Likelihood'
hdr_cells[3].text = 'Mitigation Strategy'
hdr_cells[4].text = 'Sprint Status'

risks = [
    ("Incorrect import paths or missing dependencies", "High", "Medium", "Perform code validation and testing before deployment", "Occurred S1 - Mitigated"),
    ("MongoDB connection failure", "High", "Low", "Use correct environment variables and test connectivity", "Monitored - No Issues"),
    ("Vercel deployment errors / DNS misconfiguration", "High", "High", "Monitor build logs and verify DNS 'A' records iteratively", "Occurred S2 - Active"),
    ("Data exposure (DB URI)", "High", "Low", "Store credentials in .env securely, restrict public access", "Mitigated"),
    ("Slow system response delays", "Medium", "Low", "Optimize API queries; implement Next.js caching", "Planned for S3 testing")
]

for risk, imp, like, mit, stat in risks:
    row_cells = rtable.add_row().cells
    row_cells[0].text = risk
    row_cells[1].text = imp
    row_cells[2].text = like
    row_cells[3].text = mit
    row_cells[4].text = stat

doc.save('BizMenu_Builder_HD_Revisions.docx')
